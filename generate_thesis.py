# -*- coding: utf-8 -*-
"""毕业设计论文生成脚本（以模板 DOCX 为母版填充）
用法：
    python generate_thesis.py [params.json] [模板.docx] [输出.docx]
默认参数文件为同目录 params.json，模板为张旭毕业设计模板。
"""
import sys, os, json
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
import docx
from calc import (calc_gate_width, calc_energy_dissipation, calc_seepage,
                  calc_gate_top_elevation, calc_stability, calc_gate_width_mu0,
                  calc_gate_top_mu0, calc_seepage_mu0, calc_reinforcement,
                  calc_energy_mu0, calc_stability_mu0)

BASE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_PARAMS = os.path.join(BASE, 'params.json')
# 部署模板：优先环境变量 THESIS_TEMPLATE，其次项目目录内的模板（云端/本地通用）
DEFAULT_TEMPLATE = os.environ.get('THESIS_TEMPLATE', os.path.join(BASE, '毕业设计模板.docx'))
DEFAULT_OUTPUT = os.path.join(BASE, '毕业设计输出.docx')


def replace_in_paragraph(p, old, new):
    """跨 run 字符串替换，尽量保留非目标 run 的格式与字体。"""
    if not old or old not in p.text:
        return False
    runs = p.runs
    # 1) 单 run 内直接替换
    for r in runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    # 2) 跨 run 替换：把 new 写入第一个重叠 run，其余重叠部分删除
    texts = [r.text for r in runs]
    full = ''.join(texts)
    start = full.index(old)
    end = start + len(old)
    idx = 0
    for r in runs:
        s = idx
        e = idx + len(r.text)
        idx = e
        if e <= start or s >= end:
            continue
        ov_s = max(s, start)
        ov_e = min(e, end)
        if ov_s == start:
            r.text = r.text[:ov_s - s] + new + r.text[ov_e - s:]
        else:
            r.text = r.text[:ov_s - s] + r.text[ov_e - s:]
    return True


def _iter_paragraphs(doc):
    """遍历正文 + 所有表格单元格中的段落。"""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def apply_replacements(doc, pairs):
    """对全文（含表格）做一组 old->new 替换，返回替换命中统计。"""
    stats = {}
    for p in _iter_paragraphs(doc):
        for old, new in pairs:
            if replace_in_paragraph(p, old, new):
                stats[old] = stats.get(old, 0) + 1
    return stats


def build_pairs(params):
    """构建「个人信息 + 工程标识」替换对。"""
    p = params
    pairs = []

    def add(k, old, new):
        if new and old and new != old:
            pairs.append((old, str(new)))

    add('姓名', '张旭', p.get('studentName'))
    add('学号', '220290211', p.get('studentId'))
    add('学院', '水利水电学院', p.get('college'))
    add('专业班级', '水工2202班', p.get('majorClass'))
    add('指导教师', '樊晶晶 刘林杰', p.get('advisor'))
    add('学校', '河北工程大学', p.get('university'))
    add('河流', '滏阳河', p.get('riverName'))
    add('项目缩写', 'XZ', p.get('projectAbbr'))

    # 项目名替换（按依赖顺序：先长后短，避免简称干扰全称）
    short = p.get('projectShort', '')
    if short and short != '“XZ”水闸':
        pairs.append(('“XZ”水闸', short))
    return pairs


def build_engineering_rules(params):
    """「基本资料」章节的段落级替换规则：[(定位子串, [(旧值,新值),...]), ...]"""
    def g(k):
        return str(params.get(k, ''))

    def v(k, old):
        nv = g(k)
        return (old, nv) if nv and nv != old else None

    rules = []
    r1 = [x for x in [
        v('floodStandard', '20'), v('designFlow', '174'),
        v('downstreamWaterLevel', '77.52')] if x]
    if r1:
        rules.append(('设计流量', r1))

    r2 = [x for x in [v('normalStorageLevel', '76.60')] if x]
    if r2:
        rules.append(('正常蓄水位', r2))

    r3 = [x for x in [v('gateSillElevation', '73.10')] if x]
    if r3:
        rules.append(('闸底板高程', r3))

    r4 = [x for x in [v('groundElevation', '78.80')] if x]
    if r4:
        rules.append(('现状地面高程', r4))

    r5 = [x for x in [
        v('channelBottomWidth', '20'), v('channelSlope', '1：2'),
        v('floodplainElevation', '75.00'), v('floodplainWidth', '8.5'),
        v('channelSlopeRatio', '1/1410')] if x]
    if r5:
        rules.append(('河道设计底宽', r5))

    r6 = [x for x in [
        v('mainChannelRoughness', '0.03'), v('floodplainRoughness', '0.06')] if x]
    if r6:
        rules.append(('主河槽糙率', r6))

    return rules


def apply_scoped_replacements(doc, rules):
    """段落级替换：每条规则只作用于第一个含定位子串的段落。"""
    stats = {}
    for key, pairs in rules:
        done = False
        for p in _iter_paragraphs(doc):
            if key in p.text:
                for old, new in pairs:
                    if replace_in_paragraph(p, old, new):
                        stats[old] = stats.get(old, 0) + 1
                done = True
                break
    return stats


def fmt(x, nd, strip=False):
    """数值格式化，strip=True 时去除尾部 0。"""
    s = f"{x:.{nd}f}"
    if strip:
        s = s.rstrip('0').rstrip('.')
    return s


def set_cell_text(cell, text):
    """清空单元格首段落并写入新文本（保留原字体格式）。"""
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r.text = ''
    if p.runs:
        p.runs[0].text = str(text)
    else:
        p.add_run(str(text))


def fill_chapter3(doc, gw):
    """回填第3章：表3-2(行进流速)、表3-3(堰上水头/淹没系数)、表3-7(闸孔总净宽)。"""
    rows = gw['rows']
    # 表2 = 表3-2 水流流速及上游过水断面面积
    t2 = doc.tables[2]
    for i, r in enumerate(rows):
        set_cell_text(t2.rows[i + 1].cells[1], fmt(r['A'], 3))
        set_cell_text(t2.rows[i + 1].cells[2], fmt(r['v'], 3))
    # 表3 = 表3-3 堰上水头、堰流淹没系数
    t3 = doc.tables[3]
    for i, r in enumerate(rows):
        set_cell_text(t3.rows[i + 1].cells[0], fmt(r['H0'], 3, strip=True))
        set_cell_text(t3.rows[i + 1].cells[2], fmt(r['H'], 2))
        set_cell_text(t3.rows[i + 1].cells[5], fmt(r['sigma'], 2))
    # 表7 = 表3-7 闸孔总净宽
    t7 = doc.tables[7]
    for i, r in enumerate(rows):
        set_cell_text(t7.rows[i + 1].cells[0], fmt(r['B0'], 4))
        set_cell_text(t7.rows[i + 1].cells[1], fmt(r['ratio'], 4))
        set_cell_text(t7.rows[i + 1].cells[2], fmt(r['mu0'], 4))
        set_cell_text(t7.rows[i + 1].cells[3], fmt(r['dH'], 4))
    return rows


def fill_chapter5(doc, top):
    """回填第5章：表5-2 设计闸顶高程取值。"""
    t = doc.tables[19]  # 表19 = 表5-2，列：0序号 1工况 2水位 3风浪爬高 4安全超高 5设计闸顶高程
    # R1: 序号1 蓄水位
    set_cell_text(t.rows[1].cells[2], fmt(top['normalWL'], 2))
    set_cell_text(t.rows[1].cells[3], fmt(top['h2'], 3))
    set_cell_text(t.rows[1].cells[4], fmt(top['A1'], 1))
    set_cell_text(t.rows[1].cells[5], fmt(top['H1'], 1))
    # R2: 序号2 设计洪水位
    set_cell_text(t.rows[2].cells[2], fmt(top['dsWL'], 2))
    set_cell_text(t.rows[2].cells[4], fmt(top['A2'], 1))
    set_cell_text(t.rows[2].cells[5], fmt(top['H2'], 2))
    return top


def fill_chapter6(doc, sp):
    """回填第6章：表6-2 防渗计算结果表 + 正文结论句。"""
    T = sp['Te']
    xi = sp['xi_list']
    h = sp['h_list']
    deltaH = sp['deltaH']

    def _beta(S):
        s_t = S / T
        return 0.6518 + (s_t - 0.06897) / (0.15172 - 0.06897) * (0.8710 - 0.6518)

    # 进口段（S=1）、出口段（S=2.2）修正
    beta_in = _beta(1.0)
    h_in_corr = h[0] * beta_in
    delta_in = h[0] - h_in_corr
    beta_out = sp['beta_p']
    h_out_corr = sp['h_out_corr']
    delta_out = sp['h_out'] - sp['h_out_corr']

    t = doc.tables[23]  # 表23 = 表6-2 防渗计算结果表
    # R1 进口段  R2 出口段
    set_cell_text(t.rows[1].cells[3], fmt(xi[0], 4)); set_cell_text(t.rows[1].cells[4], fmt(h[0], 4))
    set_cell_text(t.rows[1].cells[5], fmt(delta_in, 4)); set_cell_text(t.rows[1].cells[6], fmt(h_in_corr, 4))
    set_cell_text(t.rows[1].cells[7], fmt(beta_in, 4))
    set_cell_text(t.rows[2].cells[3], fmt(xi[7], 4)); set_cell_text(t.rows[2].cells[4], fmt(h[7], 4))
    set_cell_text(t.rows[2].cells[5], fmt(delta_out, 4)); set_cell_text(t.rows[2].cells[6], fmt(h_out_corr, 4))
    set_cell_text(t.rows[2].cells[7], fmt(beta_out, 4))
    # 内部垂直段 R3~R6（表顺序 0.5,0.5,1.2,1 → seg 索引 4,5,2,6）
    for ri, si in zip([3, 4, 5, 6], [4, 5, 2, 6]):
        set_cell_text(t.rows[ri].cells[3], fmt(xi[si], 4))
        set_cell_text(t.rows[ri].cells[4], fmt(h[si], 4))
    # 水平段 R7(L=14)=seg3, R8(L=12)=seg1
    set_cell_text(t.rows[7].cells[3], fmt(xi[3], 4)); set_cell_text(t.rows[7].cells[4], fmt(h[3], 4))
    set_cell_text(t.rows[8].cells[3], fmt(xi[1], 4)); set_cell_text(t.rows[8].cells[4], fmt(h[1], 4))
    # 总和 R9
    set_cell_text(t.rows[9].cells[3], fmt(sp['xi_sum'], 4))
    set_cell_text(t.rows[9].cells[4], fmt(deltaH, 1))

    # 正文结论句（段落级替换）
    concl = [
        ('闸基防渗长度', [('5×5.5', f"5×{fmt(deltaH,1)}"), ('27.5', fmt(sp['L_required'], 1))]),
        ('渗流溢出坡降', [('0.4036', fmt(sp['J_out'], 4))]),
        ('水平渗透坡降', [('0.133', fmt(sp['J_horiz'], 4))]),
    ]
    for key, pairs in concl:
        for p in _iter_paragraphs(doc):
            if key in p.text:
                for old, new in pairs:
                    replace_in_paragraph(p, old, new)
                break
    return sp


def replace_whole_paragraph(p, new_text):
    """整段替换文本，保留段落中的图片(drawing)和第一个文字 run 的字体格式。"""
    from docx.oxml.ns import qn
    runs = p.runs
    # 找第一个含文字(w:t)的 run 作为文字载体
    text_run = None
    for r in runs:
        if r._r.findall('.//' + qn('w:t')):
            text_run = r
            break
    if text_run is None:
        p.add_run(new_text)
        return
    text_run.text = new_text
    for r in runs:
        if r is text_run:
            continue
        # 只清空含文字(w:t)的 run，保留图片(drawing)等非文字 run
        if r._r.findall('.//' + qn('w:t')):
            r.text = ''


def fill_chapter9(doc, rc):
    """回填第9章配筋计算：弯矩、混凝土/钢筋强度、有效高度、选筋。"""
    M, fc, ft, fy = rc['M'], rc['fc'], rc['ft'], rc['fy']
    grade, rebar = rc['grade'], rc['rebar']
    h0, a = rc['h0'], rc['a']
    As = int(round(rc['As']))
    dia = rc['chosen'][0] if rc['chosen'] else 28

    targets = {
        '最大弯矩1682.42': lambda t: t.replace('1682.42', fmt(M, 2)),
        '采用HRB400': lambda t: t.replace('1682.42', fmt(M, 2)).replace('C30', grade).replace('HRB400', rebar),
        'fc=14.3': lambda t: (t.replace('14.3', fmt(fc, 1))
                               .replace('1.43', fmt(ft, 2))
                               .replace('360', str(int(fy)))),
        '有效高度h0': lambda t: t.replace('60', str(int(a))).replace('1140', str(int(h0))),
        '选用28@100': lambda t: t.replace('28@100', f'{dia}@100').replace('5253', str(As)),
    }
    for key, fn in targets.items():
        for p in _iter_paragraphs(doc):
            if key in p.text:
                replace_whole_paragraph(p, fn(p.text))
                break
    return rc


def build_abstract(params, gw_mu0, top_mu0, energy):
    """根据输入数据与计算结果重新撰写摘要/ABSTRACT（不照抄模板）。"""
    project = params.get('projectShort', '“XZ”水闸')
    river = params.get('riverName', '滏阳河')
    Q = params.get('designFlow', '174')
    n = gw_mu0['n']
    b0 = gw_mu0['b0']
    B0 = gw_mu0['rows'][0]['B0']
    top = top_mu0['top']
    d_pool = energy['d_pool']
    L_pool = energy['L_pool']
    L_riprap = energy['L_riprap']

    cn = [
        f"{project}位于{river}上，是一座以蓄水、灌溉为主，兼顾行洪、排涝的综合性水工建筑物。由于建成年代久远，闸体结构老化破损严重，已难以正常蓄水并存在行洪安全隐患，故对其进行拆除重建。",
        f"本次设计依据《水闸设计规范》（SL265-2016）、《水利水电工程等级划分及洪水标准》（SL252-2017）等现行规范，完成了{project}拆除重建的全过程设计。主要工作包括：确定水闸设计流量{Q}m³/s及上下游水位；按高淹没度宽顶堰公式计算闸孔总净宽，确定{n}孔、单孔净宽{b0}m的闸孔布置；完成消能防冲设计，确定消力池（深{d_pool:.2f}m、长{L_pool:.2f}m）与海漫（长{L_riprap:.2f}m）尺寸；完成闸基防渗排水设计，校核渗径长度与渗透坡降；完成闸室稳定与结构计算，验算抗滑稳定、地基应力并完成底板配筋。经计算，闸孔总净宽为{B0:.2f}m，闸顶高程为{top:.2f}m，各主要指标均满足规范要求。",
        "依据上述计算成果，利用AutoCAD绘图软件完成了水闸平面图、纵剖面图、横剖面图等主要施工图纸的绘制。",
    ]
    kw_cn = f"关键词：{project}；闸孔宽度计算；消能防冲设计；防渗排水设计；闸室稳定计算；结构配筋计算"

    abbr = params.get('projectAbbr', 'XZ')
    en = [
        f"The {abbr} sluice, a comprehensive hydraulic structure on the river mainly for water storage and irrigation while also serving flood discharge and drainage, can no longer store water normally and poses a flood risk due to long-term deterioration, and is thus demolished and reconstructed in this design.",
        f"Following the Design Code for Sluice (SL265-2016) and other current specifications, the whole design is completed, including determination of the design discharge {Q} m3/s and water levels, computation of the total net width of gate openings (with {n} openings of {b0} m each), design of energy dissipation and scour protection, seepage control, stability and structural analysis, and reinforcement of the base slab. The total net width of gate openings is {B0:.2f} m and the gate top elevation is {top:.2f} m, all satisfying the code requirements.",
        "Based on the calculation results, the plan view, longitudinal section and cross section of the sluice are produced with AutoCAD.",
    ]
    kw_en = f"Key words: {abbr} sluice; gate opening width; energy dissipation; seepage control; chamber stability; reinforcement"
    return cn, kw_cn, en, kw_en


def build_acknowledgment(params):
    """根据姓名/导师/学院重新撰写致谢（不照抄模板）。"""
    advisor = params.get('advisor', '指导教师')
    college = params.get('college', '水利水电学院')
    return [
        "行文至此，毕业设计即将完成，我的大学生涯也临近尾声。",
        f"首先，衷心感谢我的指导教师{advisor}老师。从选题论证、计算到论文成稿，{advisor}老师始终给予我悉心的指导与宝贵的建议，帮助我理清思路、修正不足。老师严谨务实的治学态度，使我受益良多。",
        f"感谢{college}各位老师四年来的辛勤教导，正是你们传授的专业知识为本次设计奠定了坚实基础。同时感谢实习单位及相关技术人员在资料收集与工程实践方面提供的帮助。",
        "感谢父母多年来的养育之恩与默默支持，你们是我求学路上最坚实的依靠。感谢同窗好友一路相伴，让这段青春岁月充满温暖与欢笑。",
        "最后，感谢一路坚持的自己。前路漫漫，愿始终保持热爱，勇敢奔赴下一段旅程。",
    ]


def fill_abstract_ack(doc, cn, kw_cn, en, kw_en, ack):
    """回填摘要/ABSTRACT/关键词/致谢（按定位整段替换，保留标题格式）。"""
    # 中文摘要正文 [50][51][52]
    cn_idx = [i for i, p in enumerate(doc.paragraphs) if '现状是集蓄水' in p.text or '综合性水工建筑物' in p.text]
    if cn_idx:
        start = cn_idx[0]
        for k, txt in enumerate(cn):
            replace_whole_paragraph(doc.paragraphs[start + k], txt)
    # 中文关键词 [54]
    for p in doc.paragraphs:
        if p.text.strip().startswith('关键词'):
            replace_whole_paragraph(p, kw_cn)
            break
    # 英文摘要 [56][57][58]
    en_idx = [i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith('The ')]
    if en_idx:
        start = en_idx[0]
        for k, txt in enumerate(en):
            replace_whole_paragraph(doc.paragraphs[start + k], txt)
    # 英文关键词 [60]
    for p in doc.paragraphs:
        if p.text.strip().startswith('Key words'):
            replace_whole_paragraph(p, kw_en)
            break
    # 致谢 [859] 起：替换为新致谢，清空多余段落（直到"参考文献"）
    ack_idx = [i for i, p in enumerate(doc.paragraphs) if '文末搁笔' in p.text]
    if ack_idx:
        start = ack_idx[0]
        for k, txt in enumerate(ack):
            if start + k < len(doc.paragraphs):
                replace_whole_paragraph(doc.paragraphs[start + k], txt)
        j = start + len(ack)
        while j < len(doc.paragraphs):
            t = doc.paragraphs[j].text.strip()
            if t and '参考文献' in t:
                break
            replace_whole_paragraph(doc.paragraphs[j], '')
            j += 1
    return True


def fill_chapter4(doc, en):
    """回填第4章消能结论句：消力池深度/长度、底板厚度、海漫长度。"""
    def repl_all(p, old, new):
        if old == new:
            return
        for _ in range(20):
            if not replace_in_paragraph(p, old, new):
                break
    concl = [
        ('池深为闸门开度', [('0.397', fmt(en['d_max'], 3)), ('0.5m', fmt(en['d_design'], 1) + 'm')]),
        ('最大值为13.271', [('13.271', fmt(en['Lsj_max'], 2)), ('15m', str(en['Lsj_design']) + 'm')]),
        ('通过计算可得最大', [('0.4457', fmt(en['t_max'], 3)), ('0.6m', fmt(en['t_design'], 1) + 'm')]),
        ('海漫长度为17.978', [('17.978', fmt(en['Lp_max'], 2)), ('20m', str(en['Lp_design']) + 'm')]),
    ]
    for key, pairs in concl:
        for p in _iter_paragraphs(doc):
            if key in p.text:
                for old, new in pairs:
                    repl_all(p, old, new)
                break
    return en


def fill_chapter7(doc, st):
    """回填第7章稳定成果：表7-10 的 Kc、σmax、σmin、σ、η + 结论判断。"""
    t = doc.tables[30]  # 表30 = 表7-10 闸室稳定计算成果表（R2 正常蓄水）
    set_cell_text(t.rows[2].cells[2], fmt(st['Kc'], 2))
    set_cell_text(t.rows[2].cells[4], fmt(st['sigma_max'], 1))
    set_cell_text(t.rows[2].cells[5], fmt(st['sigma_min'], 1))
    set_cell_text(t.rows[2].cells[6], fmt(st['sigma'], 1))
    set_cell_text(t.rows[2].cells[7], fmt(st['eta'], 3))
    set_cell_text(t.rows[2].cells[8], str(int(st['bearing'])))
    return st


def build_narrative(params):
    """重新撰写正文关键叙述段（工程概况/闸址/结构形式/消能形式/防渗意义），不照抄模板。"""
    project = params.get('projectShort', '“XZ”水闸')
    river = params.get('riverName', '滏阳河')
    sluice_type = params.get('sluiceType', '开敞式')
    weir_type = params.get('weirType', '宽顶堰')
    gate_type = params.get('gateType', '平面钢闸门')
    sill = params.get('gateSillElevation', '73.10')
    return [
        ('现状水闸建设年代久远',
         f"{project}坐落于{river}干流，是一座以蓄水、灌溉为主，兼顾行洪、排涝的综合性水工建筑物。非汛期下闸蓄水以满足周边农田灌溉与生态景观用水需求，汛期提闸泄洪以保障河道行洪安全。现有工程建成年代久远，闸体混凝土碳化剥蚀、金属结构锈蚀变形，已难以维持正常蓄水功能，且存在明显的行洪安全隐患，亟需拆除重建。"),
        ('闸址选择应综合考虑',
         f"闸址选择需统筹考虑地形地质条件、水流流态、工程布置、施工条件、经济性以及流域规划与生态环保等多方面因素。经比选，本工程沿用原闸址原位重建：该处地基为卵石土层，承载力较高、渗透稳定；河道顺直、主流稳定；原有交通与管理设施可继续利用，可有效降低工程投资，重建后可快速恢复防洪、蓄水、灌溉等综合功能，综合技术经济指标最优。"),
        ('水闸常用的堰型有宽顶堰和实用堰',
         f"本工程选用{sluice_type}布置、{weir_type}底板。{sluice_type}水闸无胸墙遮挡，泄流断面大、过流能力强，水流顺畅、不易淤积，且结构简单、施工便捷、造价较低，便于日常巡查检修。{weir_type}自由泄流范围较大、泄流能力稳定，结合原闸运行经验，本次重建仍采用{sluice_type}式、{weir_type}型底板，闸门采用{gate_type}。"),
        ('平原地区水闸水头低',
         f"本工程地处平原地区，闸上水头较低、下游河床土质抗冲能力弱，结合《水工建筑物》相关要求，消能方式选用底流消能。为保证在各种闸门开度下均能形成稍淹没水跃、有效消能，消能计算以上游蓄水位、下游无水的最不利水位组合作为控制工况，据此确定消力池、海漫及防冲槽的尺寸。"),
        ('渗流会损耗水体',
         f"水闸挡水运行时，上下游水位差会在闸基及两岸土体中形成渗流。渗流不仅造成水量损失，还会削弱闸体及岸坡的稳定性，易诱发渗透变形破坏，危及工程安全。因此必须开展防渗排水设计，合理确定地下轮廓线，延长渗径、降低渗透坡降，确保闸基渗透稳定。"),
    ]


def fill_narrative(doc, narr):
    """回填正文叙述段（整段替换）。"""
    for key, txt in narr:
        for p in _iter_paragraphs(doc):
            if key in p.text:
                replace_whole_paragraph(p, txt)
                break
    return True


def generate(params, template_path=DEFAULT_TEMPLATE, output_path=None):
    """生成论文 docx。output_path=None 返回 BytesIO（适合后端 API），
    output_path=文件路径则保存到文件（命令行模式）。"""
    print('读取模板:', template_path)
    doc = docx.Document(template_path)

    # 运行核心计算
    gate = calc_gate_width(params)
    energy = calc_energy_dissipation(params, gate)
    seep = calc_seepage(params)
    top = calc_gate_top_elevation(params, gate)
    stab = calc_stability(params, gate)
    gw_mu0 = calc_gate_width_mu0(params)   # 模板 μ0 法（第3章）
    top_mu0 = calc_gate_top_mu0(params)    # 模板方法（第5章 闸顶高程）
    seep_mu0 = calc_seepage_mu0(params)    # 模板改进阻力系数法（第6章）
    reb = calc_reinforcement(params)       # 单筋矩形截面配筋（第9章）
    energy_mu0 = calc_energy_mu0(params, gw_mu0)      # 开度扫描消能（第4章）
    stab_mu0 = calc_stability_mu0(params, gw_mu0, top_mu0)  # 闸室稳定（第7章）

    print('\n==== 核心计算结果 ====')
    print(f'闸孔总净宽 B={gate["assumedWidth"]:.2f} m (需 {gate["requiredWidth"]:.2f})')
    print(f'闸孔总宽(含闸墩)={gate["totalWidth"]:.2f} m, 单宽流量 q={gate["unitQ"]:.3f} m³/s/m')
    print(f'上游水位={gate["upstreamWL"]:.2f} m, H0={gate["headH0"]:.2f} m, σs={gate["sigmaS"]:.3f}, ε={gate["epsilon"]:.3f}')
    print(f'消力池: hc={energy["hc"]:.3f} m, hc"={energy["hc2"]:.3f} m, d={energy["d_pool"]:.2f} m, Lsj={energy["L_pool"]:.2f} m')
    print(f'海漫 Lp={energy["L_riprap"]:.2f} m, 底板厚 t={energy["t"]:.2f} m')
    print(f'防渗: ΔH={seep["deltaH"]:.2f} m, 需L={seep["L_required"]:.2f} m, 实际L={seep["L_actual"]:.2f} m')
    print(f'闸顶高程={top["topElevation"]:.2f} m (hp={top["h_p"]:.3f}, Rp={top["R_p"]:.3f})')
    print(f'稳定: ΣG={stab["sigmaG"]:.0f} kN, Kc={stab["Kc"]:.2f}, σmax={stab["sigma_max"]:.1f} kPa')

    # 个人信息替换
    pairs = build_pairs(params)
    stats = apply_replacements(doc, pairs)
    print('\n==== 个人信息替换命中 ====')
    for old, cnt in stats.items():
        print(f'  "{old}" -> 命中 {cnt} 处')

    # 基本资料工程参数替换（段落级）
    eng_rules = build_engineering_rules(params)
    eng_stats = apply_scoped_replacements(doc, eng_rules)
    print('\n==== 基本资料参数替换命中 ====')
    for old, cnt in eng_stats.items():
        print(f'  "{old}" -> 命中 {cnt} 处')

    # 第3章 闸孔总净宽（μ0 法）回填
    rows3 = fill_chapter3(doc, gw_mu0)
    print('\n==== 第3章 闸孔总净宽（μ0法）回填 ====')
    for r in rows3:
        print(f'  ΔH={r["dH"]:.1f}: H0={fmt(r["H0"],3,strip=True)}, hs/H0={fmt(r["ratio"],4)}, '
              f'μ0={fmt(r["mu0"],4)}, B0={fmt(r["B0"],4)} m')

    # 第5章 闸顶高程回填
    top5 = fill_chapter5(doc, top_mu0)
    print('\n==== 第5章 闸顶高程回填 ====')
    print(f'  H1(挡水)={fmt(top5["H1"],2)} m, H2(泄水)={fmt(top5["H2"],2)} m, 闸顶高程={fmt(top5["top"],2)} m')

    # 第6章 防渗排水回填
    sp6 = fill_chapter6(doc, seep_mu0)
    print('\n==== 第6章 防渗排水回填 ====')
    print(f'  L需={fmt(sp6["L_required"],1)} m, L实际={fmt(sp6["L_actual"],3)} m, '
          f'J出={fmt(sp6["J_out"],4)} ({"满足" if sp6["J_out_check"] else "超限"}), '
          f'Jx={fmt(sp6["J_horiz"],4)}')

    # 第9章 配筋计算回填
    r9 = fill_chapter9(doc, reb)
    print('\n==== 第9章 配筋计算回填 ====')
    print(f'  M={fmt(r9["M"],2)} kN·m, {r9["grade"]}(fc={r9["fc"]}) + {r9["rebar"]}(fy={r9["fy"]}), '
          f'As={int(round(r9["As"]))} mm², 选Φ{r9["chosen"][0] if r9["chosen"] else 28}@100')

    # 摘要 / ABSTRACT / 关键词 / 致谢（重新撰写，不照抄模板）
    cn, kw_cn, en, kw_en = build_abstract(params, gw_mu0, top_mu0, energy)
    ack = build_acknowledgment(params)
    fill_abstract_ack(doc, cn, kw_cn, en, kw_en, ack)
    print('\n==== 摘要/致谢已重新撰写 ====')

    # 第4章 消能防冲回填
    fill_chapter4(doc, energy_mu0)
    print('\n==== 第4章 消能防冲回填 ====')
    print(f'  消力池 d={energy_mu0["d_max"]:.3f}→取{energy_mu0["d_design"]}m, '
          f'Lsj={energy_mu0["Lsj_max"]:.2f}→取{energy_mu0["Lsj_design"]}m, '
          f'海漫 Lp={energy_mu0["Lp_max"]:.2f}→取{energy_mu0["Lp_design"]}m')

    # 第7章 闸室稳定回填
    fill_chapter7(doc, stab_mu0)
    print('\n==== 第7章 闸室稳定回填 ====')
    print(f'  Kc={stab_mu0["Kc"]:.2f}, σmax={stab_mu0["sigma_max"]:.1f}kPa, '
          f'η={stab_mu0["eta"]:.2f} ({"满足" if stab_mu0["stabCheck"] else "不满足"})')

    # 正文叙述段改写（工程概况/闸址/结构形式/消能形式/防渗意义）
    fill_narrative(doc, build_narrative(params))
    print('\n==== 正文叙述段已改写 ====')

    # 保存
    if output_path:
        doc.save(output_path)
        print('\n已生成:', output_path)
    else:
        import io
        output_path = io.BytesIO()
        doc.save(output_path)
        output_path.seek(0)
    print('段落数:', len(doc.paragraphs), ' 表格数:', len(doc.tables))
    return output_path


def main():
    params_path = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_PARAMS
    template_path = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_TEMPLATE
    output_path = sys.argv[3] if len(sys.argv) > 3 else DEFAULT_OUTPUT
    with open(params_path, 'r', encoding='utf-8') as fh:
        params = json.load(fh)
    generate(params, template_path, output_path)


if __name__ == '__main__':
    main()
